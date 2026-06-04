#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量修改 2027 年产品 Word 文档脚本
功能：
1. 添加缺失的出发日期
2. 更新定金金额 NZD1000
3. 修改"not exclude"为"Exclude"
4. 添加航班信息
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os
import re
from pathlib import Path

# 定义 2027 产品文件路径
BASE_PATH = r"C:\Users\Zhong\Documents\trae_projects\ChinaTravel\tours\2027"

PRODUCTS = {
    "cherry_blossoms": {
        "filename": "2027 - Cherry Blossoms 13Day.docx",
        "add_departure": "27 March 2027",
        "deposit": "NZD2500.00",  # 保持原有
        "flights": None
    },
    "china_panorama": {
        "filename": "2027 - China Panorama-27D.docx",
        "add_departure": None,
        "deposit": "NZD1000.00",
        "flights": None
    },
    "natural_china": {
        "filename": "2027 - Natural China-16D.docx",
        "add_departure": None,
        "deposit": "NZD1000.00",
        "flights": None
    },
    "silk_road": {
        "filename": "2027 - Silk Road Discovery.docx",
        "add_departure": None,
        "deposit": "NZD1000.00",
        "flights": None
    },
    "vietnam_highlight": {
        "filename": "2027 - Vietnam Highlight.docx",
        "add_departure": "12 May 2027, 15 October 2027",
        "deposit": "NZD2500.00",  # 保持原有
        "flights": None
    }
}

# 航班信息（待配置）
FLIGHT_INFO = {
    "beijing_xian": {
        "flights": [
            "CA784 on 15 October, 20:30-04:40+1",
            "CA1202 on 23 October, 17:30-20:00",
            "CA783 on 24 October, 00:05-17:30"
        ]
    },
    "chengdu_chongqing": {
        "flights": [
            "CA784 on 1 November",
            "CA1437 on 2 November, 07:00-09:50",
            "CA4117 on 9 November, 17:00-19:45",
            "CA783 on 10 November, 00:25-17:25"
        ]
    }
}

def modify_document(product_key, product_info):
    """修改单个 Word 文档"""
    filepath = os.path.join(BASE_PATH, product_info["filename"])

    if not os.path.exists(filepath):
        print(f"[ERROR] 文件不存在: {filepath}")
        return False

    try:
        doc = Document(filepath)
        modified = False

        # 遍历所有段落和表格
        for para in doc.paragraphs:
            original_text = para.text

            # 1. 修改 "not exclude" → "Exclude"
            if "not exclude" in para.text.lower():
                para.text = para.text.replace("not exclude", "Exclude")
                para.text = para.text.replace("Not exclude", "Exclude")
                modified = True
                print(f"  [OK] 修改: not exclude -> Exclude")

            # 2. 更新定金金额
            if "deposit" in para.text.lower() and "NZD" in para.text:
                # 处理定金更新
                if product_info["deposit"]:
                    para.text = re.sub(
                        r'NZD\d+,?\d*\.?\d*',
                        product_info["deposit"],
                        para.text,
                        count=1
                    )
                    modified = True
                    print(f"  [OK] 定金更新为: {product_info['deposit']}")

        # 3. 添加缺失的出发日期
        if product_info["add_departure"]:
            # 查找 Departure 或日期相关段落
            found_date_section = False
            for i, para in enumerate(doc.paragraphs):
                if "departure" in para.text.lower() or "date" in para.text.lower():
                    found_date_section = True
                    if product_info["add_departure"] not in para.text:
                        # 添加出发日期
                        para.text = para.text + f"\n{product_info['add_departure']}"
                        modified = True
                        print(f"  [OK] 添加出发日期: {product_info['add_departure']}")
                    break

        # 保存修改
        if modified:
            doc.save(filepath)
            print(f"[SUCCESS] {product_info['filename']} 已保存\n")
            return True
        else:
            print(f"[INFO] {product_info['filename']} 无需修改\n")
            return False

    except Exception as e:
        print(f"[ERROR] 处理 {product_info['filename']} 时出错: {e}\n")
        return False

def main():
    """主函数：批量修改所有产品"""
    print("=" * 60)
    print("2027 年产品批量修改脚本")
    print("=" * 60 + "\n")

    success_count = 0
    total_count = len(PRODUCTS)

    for product_key, product_info in PRODUCTS.items():
        print(f"处理: {product_info['filename']}")
        if modify_document(product_key, product_info):
            success_count += 1

    print("=" * 60)
    print(f"完成: {success_count}/{total_count} 个文件已修改")
    print("=" * 60)

if __name__ == "__main__":
    main()
